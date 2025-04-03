import pyperclip
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Expressões regulares para extrair os campos
regex_campos = {
    "Nome": r"Nome:\s*(.+)",
    "Filiação 1": r"Filiação 1:\s*(.+)",
    "Data Nascimento": r"Data Nascimento:\s*(\d{2}/\d{2}/\d{4})",
    "CPF": r"CPF:\s*([\d.-]+)",
    "Logradouro": r"Logradouro:\s*(.+)",
    "Número": r"Número:\s*(\d+)",
    "Bairro": r"Bairro:\s*(.+)",
    "CEP": r"CEP:\s*([\d-]+)",
    "Complemento": r"Complemento:\s*(.+)",
    "Cidade": r"Cidade:\s*(.+)",
    "UF": r"UF:\s*([A-Z]{2})" 
    
}

# Captura o texto da área de transferência
texto_copiado = pyperclip.paste()

def aplicar_formatacao(run):
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
    run.font.size = Pt(13)

if texto_copiado:
    dados_extraidos = {}

    # Extrai os dados do texto
    for campo, padrao in regex_campos.items():
        match = re.search(padrao, texto_copiado, re.MULTILINE)
        dados_extraidos[campo] = match.group(1) if match else "N/A"

    # Carrega o modelo .docx
    modelo = "modelo.docx"
    doc = Document(modelo)

    # Substitui os campos no corpo do texto
    for paragrafo in doc.paragraphs:
        for campo, valor in dados_extraidos.items():
            if f"{{{campo}}}" in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(f"{{{campo}}}", "")
                run = paragrafo.add_run(valor)
                aplicar_formatacao(run)

    # Substitui os campos dentro das tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for campo, valor in dados_extraidos.items():
                    if f"{{{campo}}}" in celula.text:
                        celula.text = celula.text.replace(f"{{{campo}}}", "")
                        run = celula.paragraphs[0].add_run(valor)
                        aplicar_formatacao(run)

   # Limpa o nome do arquivo para evitar caracteres inválidos
    nome_limpo = re.sub(r'[^a-zA-Z0-9_-]', '_', dados_extraidos['Nome']).strip('_')
    nome_arquivo = f"preenchido_{nome_limpo}.docx"
    
    doc.save(nome_arquivo)
    print(f"Documento preenchido com sucesso! Salvo como {nome_arquivo}")
else:
    print("Nenhum texto encontrado na área de transferência.")
input("Pressione qualquer tecla para encerrar")