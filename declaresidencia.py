import streamlit as st
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io
import base64
import os

# Configuração do Streamlit
st.set_page_config(page_title="Gerador de Declarações", layout="wide")
st.title("Gerador de Declarações")

# Expressões regulares para extrair os campos
regex_campos = {
    "Nome": r"Nome:\s*(.+)",
    "Filiação 1": r"Filiação 1:\s*(.+)",
    "Data Nascimento": r"Data Nascimento:\s*(\d{2}/\d{2}/\d{4})",
    "CPF": r"CPF:\s*([\d.-]+)",
    "Logradouro": r"Logradouro:\s*(.+)",
    "Número": r"Número:\s*(\d+)",
    "Bairro": r"Bairro:\s*(.+)",
    "CEP": r"CEP:\s*([\d-]+)",
    "Complemento": r"Complemento:\s*(.+)",
    "Cidade": r"Cidade:\s*(.+)",
    "UF": r"UF:\s*([A-Z]{2})" 
}

# Função para aplicar formatação
def aplicar_formatacao(run):
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
    run.font.size = Pt(13)

# Carrega o modelo DOCX (ajustado para Streamlit Cloud)
def load_docx_template():
    modelo_path = os.path.join(os.path.dirname(__file__), "modelo.docx")
    if os.path.exists(modelo_path):
        return Document(modelo_path)
    else:
        st.error("Arquivo modelo.docx não encontrado!")
        st.stop()

# Campo para colar o texto
texto_copiado = st.text_area("Cole o texto com os dados abaixo:", height=200)

# Botão para gerar declaração
if st.button("Gerar Declaração"):
    if texto_copiado:
        dados_extraidos = {}

        # Extrai os dados do texto
        for campo, padrao in regex_campos.items():
            match = re.search(padrao, texto_copiado, re.MULTILINE)
            dados_extraidos[campo] = match.group(1) if match else "N/A"

        try:
            doc = load_docx_template()

            # Substitui os campos no documento
            for paragrafo in doc.paragraphs:
                for campo, valor in dados_extraidos.items():
                    if f"{{{campo}}}" in paragrafo.text:
                        paragrafo.text = paragrafo.text.replace(f"{{{campo}}}", "")
                        run = paragrafo.add_run(valor)
                        aplicar_formatacao(run)

            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for campo, valor in dados_extraidos.items():
                            if f"{{{campo}}}" in celula.text:
                                celula.text = celula.text.replace(f"{{{campo}}}", "")
                                run = celula.paragraphs[0].add_run(valor)
                                aplicar_formatacao(run)

            # Gera o nome do arquivo
            nome_limpo = re.sub(r'[^a-zA-Z0-9_-]', '_', dados_extraidos.get('Nome', 'declaracao')).strip('_')
            nome_arquivo = f"declaracao_{nome_limpo}.docx"

            # Salva em memória
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            # Cria link para download
            b64 = base64.b64encode(output.read()).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{nome_arquivo}">⬇️ Baixar Declaração</a>'
            
            st.success("Declaração gerada com sucesso!")
            st.markdown(href, unsafe_allow_html=True)
            
            # Mostra os dados extraídos
            with st.expander("Ver dados extraídos"):
                st.json(dados_extraidos)

        except Exception as e:
            st.error(f"Erro ao gerar documento: {str(e)}")
    else:
        st.warning("Por favor, cole o texto com os dados antes de gerar a declaração.")

# Instruções
st.sidebar.markdown("""
### Instruções:
1. Cole o texto contendo os dados no campo ao lado
2. Clique no botão "Gerar Declaração"
3. Faça o download do documento gerado
4. Verifique os dados extraídos se necessário
""")
